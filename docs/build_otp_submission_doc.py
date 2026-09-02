from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK

OUT = 'docs/小产品实验室OTP-项目说明与创作过程.docx'
IMG_HOME = 'docs/OTP说明文档配图/01-首页预览.png'
IMG_HINT = 'docs/OTP说明文档配图/02-空状态与操作引导.png'
IMG_WORK = 'docs/OTP说明文档配图/03-网络安全运维场景.png'
IMG_BACKUP = 'docs/OTP说明文档配图/04-本地加密备份流程.png'
IMG_RESTORE = 'docs/OTP说明文档配图/05-换机恢复流程.png'
IMG_REAL_HOME = 'docs/OTP说明文档配图/06-真实截图-首页.png'
IMG_DATA = 'docs/OTP说明文档配图/08-真实截图-数据管理.png'
IMG_GENERATOR = 'docs/OTP说明文档配图/09-真实截图-密码生成器.png'
IMG_SETTINGS = 'docs/OTP说明文档配图/10-真实截图-设置页面.png'

BLUE = '2E74B5'
DARK = '17365D'
MUTED = '667085'
LIGHT = 'F2F6FC'

def set_font(run, size=11, color='1F2937', bold=None):
    # LibreOffice 的渲染环境对 Aptos 的中文回退不稳定；统一指定兼容的 Unicode 字体。
    run.font.name = 'Arial Unicode MS'
    for attr in ['ascii', 'hAnsi', 'eastAsia', 'cs']:
        run._element.rPr.rFonts.set(qn(f'w:{attr}'), 'Arial Unicode MS')
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold

def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'), 'dxa')

def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)

def set_col_widths(table, widths):
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width)

def para(doc, text='', size=11, color='1F2937', bold=False, after=6, before=0, align=None, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if align is not None: p.alignment = align
    r = p.add_run(text)
    set_font(r, size, color, bold)
    r.italic = italic
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_font(r, 16 if level == 1 else 13, BLUE if level == 1 else DARK, True)
    return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), 10.7)
    return p

def number(doc, text):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_font(p.add_run(text), 10.7)
    return p

def caption(doc, text):
    return para(doc, text, size=9.2, color=MUTED, italic=True, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)

def callout(doc, label, text):
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade(cell, LIGHT); set_cell_margins(cell, 120, 170, 120, 170)
    p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label + '  '); set_font(r, 10.6, DARK, True)
    r = p.add_run(text); set_font(r, 10.6, '344054')
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_image(doc, path, width):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(path, width=Inches(width))

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('小产品实验室 OTP  ·  ')
    set_font(run, 8.5, MUTED)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)

def main():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.78); sec.bottom_margin = Inches(0.72)
    sec.left_margin = Inches(0.82); sec.right_margin = Inches(0.82)
    sec.header_distance = Inches(0.32); sec.footer_distance = Inches(0.35)
    styles = doc.styles
    normal = styles['Normal']; normal.font.name = 'Arial Unicode MS'; normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS'); normal.font.size = Pt(11)
    for name in ['List Bullet', 'List Number']:
        styles[name].font.name = 'Arial Unicode MS'; styles[name]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial Unicode MS')
    footer = sec.footer.paragraphs[0]; add_page_number(footer)

    # Cover
    para(doc, '微信小程序项目说明文档', size=11, color=BLUE, bold=True, after=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, '小产品实验室 OTP', size=27, color=DARK, bold=True, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    para(doc, '创作背景 · 创作思路 · 实践过程 · 功能说明', size=14, color=MUTED, after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_image(doc, IMG_REAL_HOME, 3.02)
    caption(doc, '图 1  微信开发者工具中真实运行的 OTP 首页')
    callout(doc, '项目定位', '一款本地优先的 OTP 动态验证码、密码账本与加密备份工具。无需注册项目账号，用户可自行选择微信文件或自己的 WebDAV 空间保存备份。')
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(20); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run('提交材料 · 2026 年'), 9.5, MUTED)
    doc.add_page_break()

    heading(doc, '一、项目概述')
    para(doc, '小产品实验室 OTP 是一款面向微信用户的本地优先安全工具小程序，用于管理 TOTP 动态验证码、静态账号密码和加密备份。项目没有自建用户账号体系，也不要求用户注册或登录。验证码默认在设备本地计算与保存；用户可按需将加密备份导出至微信文件，或手动上传至自己配置的坚果云 WebDAV 空间，以便换机后恢复。')
    callout(doc, '核心问题', '人们在开启两步验证后，往往能顺利使用验证码，却容易在换手机、重装微信或设备遗失时，因缺少备份而无法登录重要账号。')

    heading(doc, '二、创作背景')
    heading(doc, '1. 动态验证码的“最后一公里”问题', 2)
    para(doc, '我的工作本身和网络安全运维有关，平时接触服务器、网络设备、云平台以及各种管理后台比较多。现在越来越多的平台开始强制或者建议开启二次认证，除了用户名和密码之外，还需要再输入一次动态验证码。')
    para(doc, '所以工作中，我经常会在手机里保存一些 OTP，用来登录设备、服务器或者各种网页后台。')
    callout(doc, '一个明显感受', '账号密码本身不难管理，真正容易被忽略的，是验证码换机、备份和恢复这一步。少了验证码，即使密码正确，也可能进不了服务器、云平台或管理后台。')
    add_image(doc, IMG_WORK, 6.05)
    caption(doc, '图 3  网络安全运维中的 OTP 使用场景')
    para(doc, '现有工具通常需要注册账号并把数据同步到平台云端，或者功能复杂、学习成本较高。对于只希望快速添加验证码、自己掌握数据和备份位置的用户来说，仍缺少一个轻量、直接的选择。')
    heading(doc, '2. 微信小程序的使用边界', 2)
    para(doc, '微信小程序适合“打开即用”的轻量场景，但也无法提供系统级自动填充和后台定时同步。因此，本项目不追求替代完整密码管理器，而是聚焦于 OTP 的添加、查看、密码辅助管理、备份和换机恢复。')

    heading(doc, '三、创作思路')
    heading(doc, '1. 本地优先，而不是默认上传', 2)
    para(doc, 'OTP 验证码可以根据密钥和当前时间在本地生成，因此无需在每次查看时连接服务器。项目将验证码、密码记录和备份配置优先保存在微信本地存储中；用户无需注册小程序账号，也不需要将内容上传至项目方服务器。')
    heading(doc, '2. 无需登录，也能使用自己的云端备份', 2)
    para(doc, '项目不提供自身账号登录，但允许用户连接自己的 WebDAV 空间。每次备份均遵循“本机数据 → 本地加密为备份文件 → 用户主动上传 → 自己的坚果云 WebDAV”的流程。换机时，用户从自己的备份列表中选择文件并输入备份密码恢复。')
    add_image(doc, IMG_BACKUP, 6.05)
    caption(doc, '图 4  本地加密和自主备份流程')
    callout(doc, '重要说明', '这不是后台自动同步。受微信小程序后台能力限制，备份由用户主动点击触发；项目明确告知这一边界，避免误导用户。')

    doc.add_page_break()
    heading(doc, '四、主要功能与实践实现')
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['模块', '用户能做什么', '实践方式']
    for cell, text in zip(table.rows[0].cells, headers):
        shade(cell, 'E8EEF5'); set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), 9.5, DARK, True)
    set_repeat_table_header(table.rows[0])
    rows = [
        ('OTP 验证码', '扫码、手动录入、搜索、复制验证码', '按标准 TOTP 规则在本机计算，不需要联网生成'),
        ('静态密码账本', '保存账号、密码、备注、分组；搜索与排序', '本地保存，提供手动复制，不承诺系统级自动填充'),
        ('密码生成器', '按长度、字符类型生成随机密码', '在本地生成，提供生成历史与复制提醒'),
        ('风险提示', '查看弱密码、重复密码、未备份提示', '使用本地规则判断，不读取或上传用户密码内容'),
        ('加密备份', '导出到微信文件、从微信文件恢复', '先在本地加密成备份文件，再经微信文件能力传递'),
        ('WebDAV 备份', '上传、查看并恢复自己的云端备份', '当前支持坚果云 WebDAV，用户手动发起备份'),
        ('OTP 迁移二维码', '单条或多条验证码近距离迁移', '生成加密二维码，适合换机或临时设备转移'),
        ('多语言', '跟随系统或在菜单中切换', '支持简体中文、英语、日语、繁体中文；品牌名保持中文'),
    ]
    for row_data in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_data):
            set_cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.08
            set_font(p.add_run(text), 9.1)
    set_col_widths(table, [1.12, 2.26, 3.12])
    caption(doc, '表 1  核心功能及其在小程序中的实现方式')
    add_image(doc, IMG_DATA, 2.65)
    caption(doc, '图 6  真实运行中的数据管理页面')

    heading(doc, '五、实践过程')
    headings = [
        ('阶段 1：确定功能边界', '围绕“添加 → 查看/复制 → 密码辅助管理 → 备份 → 换机恢复”建立闭环，避免堆叠与核心需求无关的复杂功能。'),
        ('阶段 2：构建本地安全流程', '将 OTP 生成、密码强度判断和备份文件生成放在本地完成。备份前由用户设置备份密码，恢复时使用同一密码解密。'),
        ('阶段 3：完善备份与迁移', '提供微信文件导入导出、坚果云 WebDAV 备份以及加密迁移二维码三条路径，对应长期留存、换机恢复和临时迁移。'),
        ('阶段 4：优化首次使用', '通过首页操作引导、清晰的三栏导航和未备份提示，帮助用户完成最重要的安全动作；语言入口置于右上角菜单，降低切换成本。'),
        ('阶段 5：功能验证', '对首页、OTP 列表与编辑、密码账本、密码生成器、回收站、设置、备份、加密迁移二维码等主要流程完成自动化回归检查，当前通过 148 项检查。'),
    ]
    for title, text in headings:
        heading(doc, title, 2); para(doc, text)

    add_image(doc, IMG_RESTORE, 6.05)
    caption(doc, '图 5  换机恢复流程：备份、导入与密码验证')

    add_image(doc, IMG_GENERATOR, 2.65)
    caption(doc, '图 7  真实运行中的密码生成器')

    doc.add_page_break()
    heading(doc, '六、交互截图与设计说明')
    add_image(doc, IMG_HINT, 2.55)
    caption(doc, '图 2  空状态中持续显示“上滑添加、下滑搜索”的操作提示')
    para(doc, '在设计上，项目把“上滑添加、下滑搜索”作为核心手势，并将提示保留在列表下方：即使用户已有内容，仍能快速回忆如何操作。首页先呈现下一步行动，而不是堆叠安全术语，降低首次使用门槛。')
    add_image(doc, IMG_SETTINGS, 2.65)
    caption(doc, '图 8  真实运行中的设置页面')
    heading(doc, '七、隐私与安全设计')
    for item in [
        '无需项目账号：不建立用户账号体系，不以注册登录作为使用前提。',
        '本地处理：验证码在设备本地计算；基础风险判断也在本地完成。',
        '用户掌握备份去向：可导出至微信文件，或手动上传到用户自己的 WebDAV。',
        '备份先加密后传递：上传或转发的是加密备份文件；恢复仍需要备份密码。',
        '不夸大能力：小程序不提供系统级自动填充，也不提供后台自动定时备份；重要账号仍建议保留服务方的备用恢复码。',
    ]: bullet(doc, item)

    heading(doc, '八、创新点与项目价值')
    for item in [
        '轻量化安全闭环：将 OTP、密码账本、备份和迁移集中在一个打开即用的小程序中。',
        '无项目账号的自主备份：不依赖项目方云端账号，用户自行决定是否连接自己的 WebDAV。',
        '双路径恢复设计：微信文件适合快速导出，WebDAV 适合稳定留存，二维码适合近距离迁移。',
        '可理解的安全提示：将未备份、弱密码等风险转化为下一步可执行操作。',
    ]: bullet(doc, item)

    heading(doc, '九、后续计划')
    para(doc, '后续将继续完善更多页面的多语言覆盖、备份状态提示和真实设备兼容性测试；同时根据用户反馈优化 WebDAV 配置引导。在不改变“本地优先、无需项目账号、用户掌握备份位置”原则的前提下，持续把安全操作做得更容易理解和完成。')
    doc.save(OUT)

if __name__ == '__main__':
    main()
